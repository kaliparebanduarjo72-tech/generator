import streamlit as st
import google.generativeai as genai
from PyPDF2 import PdfReader
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF
from io import BytesIO
import os
import re

# --- 1. KONFIGURASI API (AUTO-DETEKSI MODEL) ---
def init_api():
    try:
        api_key = st.secrets.get("GEMINI_API_KEY") or os.getenv("GEMINI_API_KEY")
        
        if not api_key:
            st.error("⚠️ API Key tidak ditemukan! Pastikan GEMINI_API_KEY sudah diset di Secrets Streamlit.")
            st.stop()
            
        genai.configure(api_key=api_key)
        
        # Mencari model yang aktif secara dinamis untuk menghindari error 404
        model_list = [
            m.name for m in genai.list_models() 
            if 'generateContent' in m.supported_generation_methods
        ]
        
        if not model_list:
            st.error("Tidak ada model AI yang tersedia untuk akun ini.")
            st.stop()
        
        # Memprioritaskan gemini-1.5-flash jika tersedia, jika tidak pakai yang pertama
        selected = "models/gemini-1.5-flash" if "models/gemini-1.5-flash" in model_list else model_list[0]
        return genai.GenerativeModel(selected)
    except Exception as e:
        st.error(f"Kesalahan Konfigurasi API: {e}")
        st.stop()

model = init_api()

# --- 2. FUNGSI EKSPOR ---
def clean_markdown_for_export(text):
    """Membersihkan sintaks markdown agar rapi di file dokumen biasa"""
    # Mengubah tabel markdown menjadi teks yang lebih terbaca
    text = text.replace("|", "  ")
    text = re.sub(r'[-]{3,}', '', text) 
    # Menghilangkan tanda bintang tebal markdown
    text = text.replace("**", "")
    return text

def export_to_word(text, school_name):
    doc = Document()
    # Header
    h = doc.add_heading(school_name, 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Isi
    cleaned_text = clean_markdown_for_export(text)
    doc.add_paragraph(cleaned_text)
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def export_to_pdf(text, school_name):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", 'B', 14)
    pdf.cell(190, 10, school_name, ln=True, align='C')
    pdf.ln(10)
    
    pdf.set_font("Arial", size=10)
    cleaned_text = clean_markdown_for_export(text)
    # Proteksi karakter encoding latin-1 agar tidak crash
    clean_text_final = cleaned_text.encode('latin-1', 'replace').decode('latin-1')
    pdf.multi_cell(0, 7, clean_text_final)
    
    return pdf.output(dest='S').encode('latin-1')

# --- 3. TAMPILAN UI ---
st.set_page_config(page_title="GuruAI - SMPN 2 Kalipare", layout="wide")

st.markdown("""
    <style>
    .header-box { background: linear-gradient(135deg, #1e3c72, #2a5298); color: white; padding: 25px; border-radius: 15px; text-align: center; }
    .main-card { background: white; padding: 25px; border-radius: 15px; box-shadow: 0 4px 10px rgba(0,0,0,0.1); margin-top: 20px; color: black; }
    
    /* CSS untuk memastikan Tabel Markdown muncul dengan garis kotak */
    div[data-testid="stMarkdownContainer"] table {
        width: 100%;
        border-collapse: collapse;
        border: 2px solid #444;
        margin: 20px 0;
    }
    div[data-testid="stMarkdownContainer"] th {
        background-color: #f2f2f2;
        padding: 10px;
        border: 1px solid #444;
        font-weight: bold;
    }
    div[data-testid="stMarkdownContainer"] td {
        padding: 10px;
        border: 1px solid #444;
    }
    </style>
    """, unsafe_allow_html=True)

st.markdown('<div class="header-box"><h1>SMP NEGERI 2 KALIPARE</h1><p>Sistem Administrasi Penilaian Otomatis</p></div>', unsafe_allow_html=True)

col1, col2 = st.columns([2, 1])

with col1:
    st.markdown('<div class="main-card">', unsafe_allow_html=True)
    input_choice = st.radio("Sumber Materi:", ["Teks Manual", "Upload PDF"])
    materi_final = ""
    if input_choice == "Teks Manual":
        materi_final = st.text_area("Masukkan Materi Pelajaran:", height=300, placeholder="Tempel teks materi di sini...")
    else:
        file_pdf = st.file_uploader("Upload PDF", type=["pdf"])
        if file_pdf:
            reader = PdfReader(file_pdf)
            for page in reader.pages: materi_final += page.extract_text()
            st.success("PDF Berhasil terbaca!")
    st.markdown('</div>', unsafe_allow_html=True)

with col2:
    st.write("⚙️ **Pengaturan Dokumen**")
    mapel = st.text_input("Mata Pelajaran", "Seni Rupa")
    bentuk_soal = st.multiselect(
        "Bentuk Soal:",
        ["Pilihan Ganda (PG)", "PG Kompleks", "Benar/Salah", "Menjodohkan", "Uraian"],
        default=["Pilihan Ganda (PG)", "Benar/Salah"]
    )
    jumlah = st.slider("Jumlah Soal", 1, 40, 5)
    
    if st.button("Generate Dokumen Rapi ✨", use_container_width=True):
        if materi_final:
            with st.spinner("AI sedang menyusun Kisi-kisi, Kartu Soal, dan Naskah..."):
                try:
                    str_bentuk = ", ".join(bentuk_soal)
                    prompt = (
                        f"Materi: {materi_final[:6000]}. Mapel: {mapel}. Sekolah: SMP NEGERI 2 KALIPARE.\n\n"
                        f"Tugas: Buatkan {jumlah} soal dengan variasi: {str_bentuk}.\n\n"
                        f"FORMAT OUTPUT (WAJIB ADA): \n"
                        f"1. KISI-KISI SOAL: (Buat Tabel Markdown dengan kolom No, Tujuan Pembelajaran, Materi, Indikator, Level, No Soal)\n"
                        f"2. KARTU SOAL: (Untuk setiap soal, buat tabel Markdown yang berisi detail indikator dan rumusan soal)\n"
                        f"3. NASKAH SOAL: (Daftar soal siap cetak)\n"
                        f"4. KUNCI JAWABAN: (Daftar jawaban benar)\n\n"
                        f"PENTING: Pastikan tabel menggunakan baris pemisah '|---|' agar tampil sebagai tabel di sistem."
                    )
                    
                    response = model.generate_content(prompt)
                    st.session_state['hasil_ujian'] = response.text
                except Exception as e:
                    st.error(f"Gagal memproses AI: {e}")
        else:
            st.warning("Mohon masukkan materi terlebih dahulu!")

# Tampilkan hasil jika data sudah ada di session
if 'hasil_ujian' in st.session_state:
    st.markdown("### 📋 Preview Dokumen")
    st.markdown(st.session_state['hasil_ujian'])
    
    st.divider()
    c1, c2 = st.columns(2)
    with c1:
        st.download_button(
            "📥 Simpan ke Word (.docx)", 
            data=export_to_word(st.session_state['hasil_ujian'], "SMP NEGERI 2 KALIPARE"), 
            file_name=f"Soal_{mapel}.docx", 
            use_container_width=True
        )
    with c2:
        st.download_button(
            "📥 Simpan ke PDF (.pdf)", 
            data=export_to_pdf(st.session_state['hasil_ujian'], "SMP NEGERI 2 KALIPARE"), 
            file_name=f"Soal_{mapel}.pdf", 
            use_container_width=True
        )
